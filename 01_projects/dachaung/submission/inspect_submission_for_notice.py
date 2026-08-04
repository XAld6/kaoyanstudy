from pathlib import Path
HERE = Path(__file__).resolve().parent


from docx import Document


DOCX = HERE / "宁夏理工学院中国国际大学生创新大赛(2026)申报书-OpenClaw版.docx"


def main() -> None:
    doc = Document(DOCX)
    text_chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_chunks.append(cell.text)
    text = "\n".join(text_chunks)

    keywords = [
        "2025年12月",
        "2026年3月",
        "高教主赛道",
        "创新训练项目",
        "一般项目",
        "结题",
        "软件著作权",
        "省级",
    ]
    for keyword in keywords:
        print(f"{keyword}: {text.count(keyword)}")

    print("--- table1 rows ---")
    table = doc.tables[0]
    for i, row in enumerate(table.rows, 1):
        seen = []
        for cell in row.cells:
            value = cell.text.replace("\n", " / ").strip()
            if value and value not in seen:
                seen.append(value)
        print(i, " || ".join(seen))


if __name__ == "__main__":
    main()

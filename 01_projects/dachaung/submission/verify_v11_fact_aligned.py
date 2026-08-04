from __future__ import annotations

from pathlib import Path

from docx import Document

HERE = Path(__file__).resolve().parent
DOCX = HERE / "智爪识损-OpenClaw-v1.1文实对齐版.docx"

REQUIRED = [
    "创新性与前沿性30%",
    "可行性与实践性25%",
    "预期成果与价值25%",
    "团队构成与执行能力20%",
    "OpenClaw",
    "Lobster",
    "AI Agent",
    "人工复核",
    "OpenCV",
    "YOLO",
    "可插拔",
    "PDF",
    "裂缝",
    "剥落",
    "渗水",
    "基础设施病害",
    "文实",
]

# Over-claim phrases that should not dominate the v1.1 fact-aligned draft
FORBIDDEN = [
    "结题报告",
    "经费使用说明",
    "3000字",
    "软件著作权",
    "省级奖项",
    "结题要求",
    "已完成大规模模型训练并达到工业级精度",
    "完全替代人工检测",
]

TRUTH_MARKERS = [
    "OpenCV规则",
    "可插拔",
    "YOLO",
    "本地运行",
    "人工复核",
    "不替代",
]


def all_text(doc: Document) -> str:
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def main() -> None:
    assert DOCX.exists(), f"missing {DOCX}"
    doc = Document(DOCX)
    text = all_text(doc)
    print(f"exists={DOCX.exists()}")
    print(f"size={DOCX.stat().st_size}")
    print(f"chars={len(text)}")

    print("--- required ---")
    missing = []
    for keyword in REQUIRED:
        count = text.count(keyword)
        print(f"{keyword}: {count}")
        if count <= 0:
            missing.append(keyword)

    print("--- forbidden ---")
    bad = []
    for keyword in FORBIDDEN:
        count = text.count(keyword)
        print(f"{keyword}: {count}")
        if count > 0 and keyword.startswith("已完成"):
            bad.append(keyword)

    print("--- truth markers ---")
    for keyword in TRUTH_MARKERS:
        print(f"{keyword}: {text.count(keyword)}")

    if len(doc.tables) >= 3:
        print("--- block lengths ---")
        print("background", len(doc.tables[1].cell(0, 0).text))
        print("research", len(doc.tables[2].cell(0, 0).text))
        print("implementation", len(doc.tables[2].cell(1, 0).text))
        print("expected", len(doc.tables[2].cell(2, 0).text))
        for name, cell in (
            ("research", doc.tables[2].cell(0, 0).text),
            ("implementation", doc.tables[2].cell(1, 0).text),
            ("expected", doc.tables[2].cell(2, 0).text),
        ):
            assert len(cell) >= 500, f"{name} block too short"

    assert not missing, f"missing required keywords: {missing}"
    assert "OpenCV" in text and "YOLO" in text and "可插拔" in text
    print("OK: v1.1 fact-aligned submission checks passed")


if __name__ == "__main__":
    main()

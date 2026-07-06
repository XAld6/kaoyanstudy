from pathlib import Path

from docx import Document


DOCX = Path(r"D:\xm\dachaung\智爪识损-基于OpenClaw的大创通知合规扩写版.docx")


REQUIRED_KEYWORDS = [
    "创新训练项目",
    "一般项目",
    "2026年6月——2027年6月",
    "1年",
    "兴趣驱动、自主实践、重在过程",
    "创新性",
    "可行性",
    "预期成果",
    "团队执行能力",
    "结题报告",
    "经费使用说明",
    "软件著作权",
    "不少于3000字",
    "省级奖项",
    "OpenClaw",
    "Lobster",
    "人工复核",
    "乡村",
    "社区",
]

FORBIDDEN = [
    "高教主赛道",
    "青年红色筑梦之旅赛道",
    "职教赛道",
    "产业命题赛道",
    "2025年12月—— 2026年3月",
]


def all_text(doc: Document) -> str:
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def print_table_one(doc: Document) -> None:
    table = doc.tables[0]
    print("--- table1 rows ---")
    for i, row in enumerate(table.rows[:15], 1):
        seen = []
        for cell in row.cells:
            value = cell.text.replace("\n", " / ").strip()
            if value and value not in seen:
                seen.append(value)
        print(i, " || ".join(seen))


def main() -> None:
    doc = Document(DOCX)
    text = all_text(doc)
    print(f"exists={DOCX.exists()}")
    print(f"size={DOCX.stat().st_size}")
    print(f"chars={len(text)}")
    print("--- required keyword counts ---")
    for keyword in REQUIRED_KEYWORDS:
        print(f"{keyword}: {text.count(keyword)}")
    print("--- forbidden keyword counts ---")
    for keyword in FORBIDDEN:
        print(f"{keyword}: {text.count(keyword)}")
    if len(doc.tables) >= 3:
        print("--- block lengths ---")
        print("background", len(doc.tables[1].cell(0, 0).text))
        print("research", len(doc.tables[2].cell(0, 0).text))
        print("implementation", len(doc.tables[2].cell(1, 0).text))
        print("expected", len(doc.tables[2].cell(2, 0).text))
    print_table_one(doc)


if __name__ == "__main__":
    main()

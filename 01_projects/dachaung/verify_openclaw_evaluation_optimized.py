from pathlib import Path

from docx import Document


DOCX = Path(r"D:\xm\dachaung\智爪识损-OpenClaw评审优化丰富版.docx")

REQUIRED = [
    "创新性与前沿性30%",
    "可行性与实践性25%",
    "预期成果与价值25%",
    "团队构成与执行能力20%",
    "OpenClaw",
    "Lobster",
    "AI Agent",
    "人工复核",
    "历史追踪",
    "基础设施病害",
]

FORBIDDEN = [
    "结题报告",
    "经费使用说明",
    "3000字",
    "不少于3000字",
    "软件著作权",
    "省级奖项",
    "结题要求",
]


def all_text(doc: Document) -> str:
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def main() -> None:
    doc = Document(DOCX)
    text = all_text(doc)
    print(f"exists={DOCX.exists()}")
    print(f"size={DOCX.stat().st_size}")
    print(f"chars={len(text)}")
    print("--- required ---")
    for keyword in REQUIRED:
        print(f"{keyword}: {text.count(keyword)}")
    print("--- forbidden ---")
    for keyword in FORBIDDEN:
        print(f"{keyword}: {text.count(keyword)}")
    if len(doc.tables) >= 3:
        print("--- block lengths ---")
        print("background", len(doc.tables[1].cell(0, 0).text))
        print("research", len(doc.tables[2].cell(0, 0).text))
        print("implementation", len(doc.tables[2].cell(1, 0).text))
        print("expected", len(doc.tables[2].cell(2, 0).text))


if __name__ == "__main__":
    main()

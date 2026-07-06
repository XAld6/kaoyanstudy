from pathlib import Path

from docx import Document


PATH = Path("D:/xm/dachaung/宁夏理工学院中国国际大学生创新大赛(2026)申报书-OpenClaw版.docx")
OLD_NAME = "慧眼识裂——基础设施安全智能诊断系统"
NEW_NAME = "智爪识损——基于OpenClaw的基础设施病害AI全自动识别与诊断系统"
KEYWORDS = ["OpenClaw", "Lobster", "AI Agent", "工作流", "人工复核"]


def all_text(doc: Document) -> str:
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def main() -> None:
    doc = Document(PATH)
    text = all_text(doc)
    print(f"exists={PATH.exists()}")
    print(f"size={PATH.stat().st_size}")
    print(f"old_name_count={text.count(OLD_NAME)}")
    print(f"new_name_count={text.count(NEW_NAME)}")
    for keyword in KEYWORDS:
        print(f"{keyword}_count={text.count(keyword)}")
    if len(doc.tables) >= 3:
        print(f"background_chars={len(doc.tables[1].cell(0, 0).text)}")
        print(f"research_chars={len(doc.tables[2].cell(0, 0).text)}")
        print(f"implementation_chars={len(doc.tables[2].cell(1, 0).text)}")
        print(f"expected_chars={len(doc.tables[2].cell(2, 0).text)}")


if __name__ == "__main__":
    main()
